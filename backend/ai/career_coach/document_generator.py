import io
import datetime
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Dict, Any, List

def generate_cover_letter_docx(cover_letter_content: str, candidate_info: Dict[str, Any], job_info: Dict[str, Any] = None) -> bytes:
    """
    Generates a professionally formatted DOCX cover letter.
    """
    doc = Document()
    
    # Set standard margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Header block
    name = candidate_info.get('name', 'Candidate Name')
    email = candidate_info.get('email', '')
    phone = candidate_info.get('phone', '')
    location = candidate_info.get('location', '')
    
    # Add candidate name in bold/large font
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    run_name = title_p.add_run(name)
    run_name.font.size = Pt(18)
    run_name.bold = True
    run_name.font.color.rgb = RGBColor(31, 78, 121) # Sleek blue
    
    # Contact details subheader
    contact_p = doc.add_paragraph()
    contact_p.paragraph_format.space_after = Pt(24)
    contact_items = [item for item in [email, phone, location] if item]
    contact_text = " | ".join(contact_items)
    run_contact = contact_p.add_run(contact_text)
    run_contact.font.size = Pt(9.5)
    run_contact.font.color.rgb = RGBColor(100, 100, 100) # Gray
    
    # Add Date
    date_p = doc.add_paragraph()
    date_p.paragraph_format.space_after = Pt(12)
    date_p.add_run(datetime.date.today().strftime("%B %d, %Y"))
    
    # Add Recruiter/Company Address Block
    if job_info:
        company = job_info.get('company', '')
        title = job_info.get('title', '')
        
        addr_p = doc.add_paragraph()
        addr_p.paragraph_format.space_after = Pt(18)
        addr_p.paragraph_format.line_spacing = 1.15
        
        run_to = addr_p.add_run("Hiring Team\n")
        run_to.bold = True
        
        if company:
            addr_p.add_run(f"{company}\n")
        if title:
            addr_p.add_run(f"Re: {title} Position\n")
            
    # Body Paragraphs
    paragraphs = cover_letter_content.split('\n')
    for p_text in paragraphs:
        p_text_stripped = p_text.strip()
        if not p_text_stripped:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(p_text_stripped)
        
    # Save document to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()

def generate_optimized_resume_docx(parsed_data: Dict[str, Any], optimizations: Dict[str, Any] = None) -> bytes:
    """
    Generates a formatted resume from parsed and AI-optimized data.
    """
    doc = Document()
    
    # Margins (0.75 inch)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    
    # 1. Header (Name, contact details)
    personal = parsed_data.get('personal_info', {}) or {}
    name = personal.get('name', parsed_data.get('name', 'Candidate Name'))
    email = personal.get('email', parsed_data.get('email', ''))
    phone = personal.get('phone', parsed_data.get('phone', ''))
    location = personal.get('location', parsed_data.get('location', ''))
    
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.paragraph_format.space_after = Pt(4)
    run_name = header_p.add_run(name)
    run_name.font.size = Pt(20)
    run_name.bold = True
    run_name.font.color.rgb = RGBColor(31, 78, 121)
    
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(18)
    contact_items = [item for item in [email, phone, location] if item]
    contact_text = " | ".join(contact_items)
    run_contact = contact_p.add_run(contact_text)
    run_contact.font.size = Pt(9.5)
    run_contact.font.color.rgb = RGBColor(100, 100, 100)
    
    # Helper to add section headers
    def add_section_header(title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(title_text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(31, 78, 121)
        
        # Add a thin bottom border under the header
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F4E79')
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)
        
    # 2. Professional Summary
    summary = optimizations.get('summary', parsed_data.get('summary', '')) if optimizations else parsed_data.get('summary', '')
    if summary:
        add_section_header("Professional Summary")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(summary)
        
    # 3. Skills
    skills = optimizations.get('skills', parsed_data.get('skills', [])) if optimizations else parsed_data.get('skills', [])
    if skills:
        add_section_header("Key Skills")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.15
        
        skill_names = []
        for s in skills:
            if isinstance(s, dict):
                skill_names.append(s.get('name', ''))
            else:
                skill_names.append(str(s))
        skill_names = [s for s in skill_names if s]
        
        run_skills = p.add_run(", ".join(skill_names))
        run_skills.font.size = Pt(10)
        
    # 4. Work Experience
    experience = optimizations.get('experience', parsed_data.get('experience', [])) if optimizations else parsed_data.get('experience', [])
    if experience:
        add_section_header("Professional Experience")
        for idx, exp in enumerate(experience):
            title = exp.get('title', '')
            company = exp.get('company', '')
            start = exp.get('start_date', '') or exp.get('startDate', '')
            end = exp.get('end_date', '') or exp.get('endDate', '') or 'Present'
            desc = exp.get('description', '')
            bullets = exp.get('bullets', [])
            
            # Title & Company
            p_title = doc.add_paragraph()
            p_title.paragraph_format.space_before = Pt(6)
            p_title.paragraph_format.space_after = Pt(2)
            p_title.paragraph_format.keep_with_next = True
            
            run_title = p_title.add_run(f"{title}")
            run_title.bold = True
            
            if company:
                run_company = p_title.add_run(f" | {company}")
                run_company.italic = True
                
            if start:
                run_dates = p_title.add_run(f"\t{start} – {end}")
                # Right align dates
                tab_stops = p_title.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(7.0), alignment=2)
                
            # If there's a text description
            if desc:
                p_desc = doc.add_paragraph()
                p_desc.paragraph_format.space_after = Pt(4)
                p_desc.paragraph_format.line_spacing = 1.15
                p_desc.add_run(desc)
                
            # Experience Bullet Points
            if bullets:
                for bullet in bullets:
                    if not bullet.strip():
                        continue
                    p_bullet = doc.add_paragraph(style='List Bullet')
                    p_bullet.paragraph_format.space_after = Pt(3)
                    p_bullet.paragraph_format.line_spacing = 1.15
                    bullet_cleaned = re.sub(r'^[•\-\*\s]+', '', bullet)
                    p_bullet.add_run(bullet_cleaned)
            elif desc and '\n' in desc:
                # If bullets are embedded as newlines
                lines = desc.split('\n')
                p_desc.clear()
                for line in lines:
                    line_clean = line.strip()
                    if not line_clean:
                        continue
                    if line_clean.startswith(('•', '-', '*')):
                        p_b = doc.add_paragraph(style='List Bullet')
                        p_b.paragraph_format.space_after = Pt(3)
                        p_b.paragraph_format.line_spacing = 1.15
                        p_b.add_run(re.sub(r'^[•\-\*\s]+', '', line_clean))
                    else:
                        p_l = doc.add_paragraph()
                        p_l.paragraph_format.space_after = Pt(4)
                        p_l.paragraph_format.line_spacing = 1.15
                        p_l.add_run(line_clean)
                        
    # 5. Education
    education = optimizations.get('education', parsed_data.get('education', [])) if optimizations else parsed_data.get('education', [])
    if education:
        add_section_header("Education")
        for edu in education:
            degree = edu.get('degree', '')
            field = edu.get('field', '') or edu.get('major', '')
            inst = edu.get('institution', '') or edu.get('school', '')
            year = edu.get('graduation_year', '') or edu.get('year', '')
            
            p_edu = doc.add_paragraph()
            p_edu.paragraph_format.space_before = Pt(4)
            p_edu.paragraph_format.space_after = Pt(2)
            
            edu_title = f"{degree}"
            if field:
                edu_title += f" in {field}"
            run_edu = p_edu.add_run(edu_title)
            run_edu.bold = True
            
            if inst:
                run_inst = p_edu.add_run(f" | {inst}")
                run_inst.italic = True
                
            if year:
                p_edu.add_run(f"\t{year}")
                tab_stops = p_edu.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(7.0), alignment=2)
                
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()
